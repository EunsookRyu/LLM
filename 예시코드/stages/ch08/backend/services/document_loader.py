# services/document_loader.py
# 가이드 참조: chapter8_document_indexing.md L33–L235
import io
from pathlib import Path
from dataclasses import dataclass, field


@dataclass
class DocumentPage:
    """
    문서에서 추출한 페이지 단위의 텍스트와 메타데이터를 담는 데이터 클래스.
    청킹 이전의 원본 단위를 표현한다.
    """
    text: str
    metadata: dict = field(default_factory=dict)


class DocumentLoader:
    """
    다양한 파일 형식에서 텍스트를 추출하는 로더 클래스.
    파일 확장자를 기반으로 적절한 처리 방법을 자동으로 선택한다.
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}

    def load(self, file_path: str | Path) -> list[DocumentPage]:
        """
        파일 경로를 받아 텍스트를 추출한다.
        파일 형식에 따라 적절한 로더를 자동으로 선택한다.
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"파일을 찾을 수 없습니다: {file_path}")

        extension = path.suffix.lower()

        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"지원하지 않는 파일 형식입니다: {extension}\n"
                f"지원 형식: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )

        loader_map = {
            ".pdf": self._load_pdf,
            ".docx": self._load_docx,
            ".txt": self._load_text,
            ".md": self._load_text,
        }

        return loader_map[extension](path)

    def load_from_bytes(
        self,
        file_bytes: bytes,
        filename: str
    ) -> list[DocumentPage]:
        """
        파일 경로 대신 바이트 데이터로 직접 문서를 불러온다.
        API를 통해 파일을 업로드받을 때 사용한다.
        """
        extension = Path(filename).suffix.lower()

        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"지원하지 않는 파일 형식입니다: {extension}"
            )

        if extension == ".pdf":
            return self._load_pdf_from_bytes(file_bytes, filename)
        elif extension == ".docx":
            return self._load_docx_from_bytes(file_bytes, filename)
        else:
            text = file_bytes.decode("utf-8", errors="replace")
            return [DocumentPage(
                text=text,
                metadata={"source": filename, "page": 1}
            )]

    # ─────────────────────────────────────────────
    # PDF 로더
    # ─────────────────────────────────────────────

    def _load_pdf(self, path: Path) -> list[DocumentPage]:
        with open(path, "rb") as f:
            return self._load_pdf_from_bytes(f.read(), path.name)

    def _load_pdf_from_bytes(
        self,
        file_bytes: bytes,
        filename: str
    ) -> list[DocumentPage]:
        import fitz  # pymupdf

        pages = []
        doc = fitz.open(stream=file_bytes, filetype="pdf")

        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()

            # 텍스트가 없는 페이지(이미지만 있는 경우 등)는 건너뛴다.
            text = text.strip()
            if not text:
                continue

            pages.append(DocumentPage(
                text=text,
                metadata={
                    "source": filename,
                    "page": page_num + 1,
                    "total_pages": len(doc),
                }
            ))

        doc.close()

        if not pages:
            raise ValueError(
                f"'{filename}'에서 텍스트를 추출할 수 없습니다. "
                "스캔 이미지 PDF이거나 텍스트가 없는 파일일 수 있습니다."
            )

        return pages

    # ─────────────────────────────────────────────
    # DOCX 로더
    # ─────────────────────────────────────────────

    def _load_docx(self, path: Path) -> list[DocumentPage]:
        with open(path, "rb") as f:
            return self._load_docx_from_bytes(f.read(), path.name)

    def _load_docx_from_bytes(
        self,
        file_bytes: bytes,
        filename: str
    ) -> list[DocumentPage]:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        full_text = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                full_text.append(text)

        # 표 안의 텍스트도 추출한다.
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        full_text.append(text)

        if not full_text:
            raise ValueError(
                f"'{filename}'에서 텍스트를 추출할 수 없습니다."
            )

        # DOCX는 페이지 구분이 불명확하므로 전체를 하나의 페이지로 취급한다.
        return [DocumentPage(
            text="\n".join(full_text),
            metadata={
                "source": filename,
                "page": 1,
                "total_pages": 1,
            }
        )]

    # ─────────────────────────────────────────────
    # 텍스트 로더 (TXT, MD)
    # ─────────────────────────────────────────────

    def _load_text(self, path: Path) -> list[DocumentPage]:
        # 인코딩을 자동으로 감지하여 읽는다.
        # UTF-8을 먼저 시도하고 실패하면 CP949(EUC-KR)로 재시도한다.
        for encoding in ["utf-8", "cp949", "euc-kr"]:
            try:
                text = path.read_text(encoding=encoding)
                break
            except UnicodeDecodeError:
                continue
        else:
            raise ValueError(
                f"'{path.name}'의 인코딩을 감지하지 못했습니다. "
                "UTF-8 또는 EUC-KR로 저장된 파일인지 확인하세요."
            )

        return [DocumentPage(
            text=text,
            metadata={
                "source": path.name,
                "page": 1,
                "total_pages": 1,
            }
        )]


# 싱글톤 인스턴스
document_loader = DocumentLoader()
